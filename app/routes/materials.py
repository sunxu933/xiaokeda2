"""Materials routes - Learning resources management."""
import os
import json
import uuid
import re
from datetime import datetime
from pathlib import Path
from flask import Blueprint, render_template, request, redirect, url_for, flash, g, jsonify, send_file
from werkzeug.utils import secure_filename

from app.extensions import db
from app.models.material import Material, MaterialKnowledge, MaterialQuestion
from app.models.student import Student
from app.models.settings import AppSetting
from app.services.ai_service import ai_service
from app.services.pdf_service import pdf_service
from app.helpers import save_uploaded_file, infer_subject_from_filename, infer_grade_from_filename

materials_bp = Blueprint('materials', __name__)


@materials_bp.route('/')
def list():
    """List all materials."""
    student = g.current_student
    if not student:
        return redirect(url_for('settings.student_profile'))

    subject = request.args.get('subject')
    material_type = request.args.get('type')
    status = request.args.get('status')

    query = Material.query.filter_by(student_id=student.id)

    if subject:
        query = query.filter_by(subject=subject)
    if material_type:
        query = query.filter_by(material_type=material_type)
    if status:
        query = query.filter_by(status=status)

    materials = query.order_by(Material.created_at.desc()).all()

    return render_template('materials/list.html', materials=materials)


@materials_bp.route('/upload', methods=['GET', 'POST'])
def upload():
    """Upload new materials."""
    student = g.current_student
    if not student:
        return redirect(url_for('settings.student_profile'))

    if request.method == 'POST':
        files = request.files.getlist('files')
        auto_analyze = request.form.get('auto_analyze') == 'on'

        for file in files:
            if file.filename == '':
                continue

            ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
            filepath = save_uploaded_file(file, 'materials')

            # Infer metadata from filename
            subject = infer_subject_from_filename(file.filename)
            grade = infer_grade_from_filename(file.filename)
            material_type = None
            if '试卷' in file.filename or '考试' in file.filename:
                material_type = '试卷'
            elif '练习' in file.filename:
                material_type = '练习册'
            elif '讲义' in file.filename:
                material_type = '讲义'
            elif '课本' in file.filename or '教材' in file.filename:
                material_type = '课本'

            material = Material(
                student_id=student.id,
                title=secure_filename(file.filename),
                file_path=filepath,
                file_type=ext,
                subject=subject,
                source_grade=grade or student.grade,
                material_type=material_type,
                status='pending' if auto_analyze else 'pending'
            )
            db.session.add(material)
            db.session.flush()

            # Auto analyze if requested
            if auto_analyze:
                _analyze_material_async(material, student)

        db.session.commit()
        flash(f'已上传 {len(files)} 个文件', 'success')
        return redirect(url_for('materials.list'))

    return render_template('materials/upload.html')


@materials_bp.route('/<int:id>')
def detail(id):
    """View material detail."""
    student = g.current_student
    material = Material.query.get_or_404(id)

    if material.student_id != student.id:
        return redirect(url_for('materials.list'))

    return render_template('materials/detail.html', material=material)


@materials_bp.route('/<int:id>/preview')
def preview(id):
    """Preview material file."""
    student = g.current_student
    material = Material.query.get_or_404(id)

    if material.student_id != student.id:
        return redirect(url_for('materials.list'))

    if material.file_type == 'image':
        return render_template('materials/preview.html', material=material)
    else:
        return redirect(url_for('materials.detail', id=id))


@materials_bp.route('/<int:id>/analyze', methods=['POST'])
def analyze(id):
    """Trigger AI analysis of material."""
    student = g.current_student
    material = Material.query.get_or_404(id)

    if material.student_id != student.id:
        return redirect(url_for('materials.list'))

    try:
        _analyze_material(material, student)
        flash('分析完成', 'success')
    except Exception as e:
        flash(f'分析失败: {str(e)}', 'error')
        material.status = 'failed'

    db.session.commit()
    return redirect(url_for('materials.detail', id=id))


@materials_bp.route('/<int:id>/delete', methods=['POST'])
def delete(id):
    """Delete material."""
    student = g.current_student
    material = Material.query.get_or_404(id)

    if material.student_id != student.id:
        return redirect(url_for('materials.list'))

    # Delete file
    if material.file_path:
        full_path = Path(__file__).parent.parent / 'app' / 'static' / material.file_path
        if full_path.exists():
            full_path.unlink()

    db.session.delete(material)
    db.session.commit()
    flash('资料已删除', 'success')
    return redirect(url_for('materials.list'))


@materials_bp.route('/knowledge-base')
def knowledge_base():
    """Browse knowledge base."""
    student = g.current_student
    if not student:
        return redirect(url_for('settings.student_profile'))

    from app.models.knowledge import KnowledgePoint

    subject = request.args.get('subject')
    grade = request.args.get('grade', type=int) or student.grade

    query = KnowledgePoint.query.filter_by(grade=grade)
    if subject:
        query = query.filter_by(subject=subject)

    points = query.order_by(KnowledgePoint.subject, KnowledgePoint.semester, KnowledgePoint.chapter).all()

    # Group by chapter
    chapters = {}
    for p in points:
        key = (p.subject, p.semester, p.chapter)
        if key not in chapters:
            chapters[key] = {
                'subject': p.subject,
                'semester': p.semester,
                'chapter': p.chapter,
                'topics': []
            }
        chapters[key]['topics'].append(p)

    return render_template('materials/knowledge_base.html', chapters=list(chapters.values()), student=student)


@materials_bp.route('/knowledge-base/search')
def search_knowledge():
    """Search knowledge base (JSON)."""
    student = g.current_student
    if not student:
        return jsonify({'error': 'No student'}), 400

    query = request.args.get('q', '')
    subject = request.args.get('subject')

    from app.models.knowledge import KnowledgePoint

    q = KnowledgePoint.query
    if subject:
        q = q.filter_by(subject=subject)

    if query:
        q = q.filter(
            db.or_(
                KnowledgePoint.topic.ilike(f'%{query}%'),
                KnowledgePoint.chapter.ilike(f'%{query}%')
            )
        )

    results = q.limit(20).all()

    return jsonify({
        'results': [{
            'id': p.id,
            'subject': p.subject,
            'grade': p.grade,
            'chapter': p.chapter,
            'topic': p.topic
        } for p in results]
    })


@materials_bp.route('/generate-questions', methods=['POST'])
def generate_questions():
    """Generate practice questions from knowledge points (JSON)."""
    student = g.current_student
    if not student:
        return jsonify({'error': 'No student'}), 400

    data = request.get_json()
    knowledge_points = data.get('knowledge_points', [])
    subject = data.get('subject', '数学')
    count = data.get('count', 5)

    try:
        result = ai_service.generate_questions_from_knowledge(knowledge_points, subject, student.grade, count)
        return jsonify({'success': True, 'questions': result.get('questions', [])})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@materials_bp.route('/browse')
def browse():
    """Browse local directory for materials."""
    student = g.current_student
    if not student:
        return redirect(url_for('settings.student_profile'))

    base_dir = AppSetting.get('materials_local_dir', '')
    if not base_dir:
        return render_template('materials/browse.html', files=[], base_dir='')

    return render_template('materials/browse.html', files=[], base_dir=base_dir)


@materials_bp.route('/browse/api/list')
def browse_list():
    """List files in local directory (JSON)."""
    base_dir = AppSetting.get('materials_local_dir', '')
    if not base_dir:
        return jsonify({'error': 'No base directory configured'}), 400

    path = request.args.get('path', '')

    # Security: prevent path traversal using Path.resolve()
    try:
        base_path = Path(base_dir).resolve()
        if not base_path.exists():
            return jsonify({'error': 'Base directory does not exist'}), 400
        if not base_path.is_dir():
            return jsonify({'error': 'Base path is not a directory'}), 400

        full_path = (base_path / path).resolve()

        # Ensure resolved path is within base directory
        if not str(full_path).startswith(str(base_path)):
            return jsonify({'error': 'Invalid path'}), 400

        if not full_path.is_dir():
            return jsonify({'error': 'Not a directory'}), 400
    except (OSError, ValueError):
        return jsonify({'error': 'Invalid path'}), 400

    items = []
    for item in full_path.iterdir():
        rel_path = item.relative_to(base_path)
        items.append({
            'name': item.name,
            'path': str(rel_path),
            'is_dir': item.is_dir()
        })

    items.sort(key=lambda x: (not x['is_dir'], x['name']))

    return jsonify({
        'items': items,
        'parent': str(Path(path).parent) if path else None
    })


@materials_bp.route('/browse/api/import', methods=['POST'])
def browse_import():
    """Import file from local directory (JSON)."""
    student = g.current_student
    if not student:
        return jsonify({'error': 'No student'}), 400

    data = request.get_json()
    source_path = data.get('path', '').lstrip('/')

    base_dir = AppSetting.get('materials_local_dir', '')
    if not base_dir:
        return jsonify({'error': 'No base directory configured'}), 400

    # Security: use Path.resolve() to prevent path traversal
    try:
        base_path = Path(base_dir).resolve()
        if not base_path.exists() or not base_path.is_dir():
            return jsonify({'error': 'Invalid base directory'}), 400

        full_source = (base_path / source_path).resolve()

        if not str(full_source).startswith(str(base_path)):
            return jsonify({'error': 'Invalid path'}), 400

        if not full_source.exists() or not full_source.is_file():
            return jsonify({'error': 'File not found'}), 400
    except (OSError, ValueError):
        return jsonify({'error': 'Invalid path'}), 400

    # Copy file
    import shutil
    filename = full_source.name
    dest_subfolder = 'materials'
    dest_path = Path(__file__).parent.parent / 'app' / 'static' / 'uploads' / dest_subfolder
    dest_path.mkdir(parents=True, exist_ok=True)

    ext = full_source.suffix.lstrip('.').lower() if '.' in filename else ''
    unique_name = f"{uuid.uuid4().hex}.{ext}"
    dest_file = dest_path / unique_name

    shutil.copy2(full_source, dest_file)

    # Create material record
    subject = infer_subject_from_filename(filename)
    grade = infer_grade_from_filename(filename)

    material = Material(
        student_id=student.id,
        title=filename,
        file_path=f"uploads/{dest_subfolder}/{unique_name}",
        file_type=ext,
        subject=subject,
        source_grade=grade or student.grade,
        status='pending'
    )
    db.session.add(material)
    db.session.commit()

    return jsonify({'success': True, 'material_id': material.id})


@materials_bp.route('/browse/api/set-base-dir', methods=['POST'])
def set_base_dir():
    """Set local browsing base directory (JSON)."""
    data = request.get_json()
    path = data.get('path', '')

    AppSetting.set('materials_local_dir', path)
    return jsonify({'success': True})


@materials_bp.route('/compile')
def compile():
    """Compile questions from multiple materials into PDF."""
    student = g.current_student
    if not student:
        return redirect(url_for('settings.student_profile'))

    materials = Material.query.filter_by(
        student_id=student.id,
        status='analyzed'
    ).all()

    return render_template('materials/compile.html', materials=materials)


@materials_bp.route('/compile/api/types', methods=['POST'])
def compile_types():
    """Get available question types from selected materials (JSON)."""
    data = request.get_json()
    material_ids = data.get('material_ids', [])

    questions = MaterialQuestion.query.filter(
        MaterialQuestion.material_id.in_(material_ids)
    ).all()

    types = set(q.question_type for q in questions if q.question_type)

    return jsonify({'types': list(types)})


@materials_bp.route('/compile/api/generate', methods=['POST'])
def compile_generate():
    """Generate compiled PDF (JSON)."""
    student = g.current_student
    if not student:
        return jsonify({'error': 'No student'}), 400

    data = request.get_json()
    material_ids = data.get('material_ids', [])
    question_types = data.get('question_types', [])
    title = data.get('title', '专项练习')

    # Query questions
    query = MaterialQuestion.query.filter(MaterialQuestion.material_id.in_(material_ids))
    if question_types:
        query = query.filter(MaterialQuestion.question_type.in_(question_types))

    questions = query.all()

    # Group by material
    sections = []
    for m_id in material_ids:
        material = db.session.get(Material, m_id)
        if not material:
            continue

        m_questions = [q for q in questions if q.material_id == m_id]
        if m_questions:
            sections.append({
                'material_title': material.title,
                'material': material,
                'questions': [q.to_dict() for q in m_questions]
            })

    # Generate PDF
    try:
        pdf_content = pdf_service.build(title, sections, include_answers=True)
        filepath = pdf_service.save_compiled_pdf(pdf_content)

        # Cleanup old PDFs
        pdf_service.cleanup_old_pdfs(hours=1)

        return jsonify({'success': True, 'filepath': filepath})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@materials_bp.route('/compile/preview/<path:filename>')
def compile_preview(filename):
    """Preview compiled PDF."""
    student = g.current_student
    if not student:
        return redirect(url_for('settings.student_profile'))

    filepath = Path(__file__).parent.parent / 'app' / 'static' / 'uploads' / 'materials' / filename
    if not filepath.exists():
        flash('文件不存在', 'error')
        return redirect(url_for('materials.compile'))

    return send_file(filepath, mimetype='application/pdf')


@materials_bp.route('/practice')
def practice():
    """Practice page using knowledge base."""
    student = g.current_student
    if not student:
        return redirect(url_for('settings.student_profile'))

    return render_template('materials/practice.html', student=student)


# === Helper Functions ===

def _analyze_material(material, student):
    """Analyze a material using AI."""
    material.status = 'analyzing'
    db.session.commit()

    try:
        if material.file_type == 'image':
            _analyze_image_material(material, student)
        elif material.file_type == 'pdf':
            _analyze_pdf_material(material, student)
        elif material.file_type == 'docx':
            _analyze_docx_material(material, student)
        else:
            raise ValueError(f"Unsupported file type: {material.file_type}")

        # Infer metadata from title
        _infer_metadata_from_title(material)

        material.status = 'analyzed'
    except Exception as e:
        material.status = 'failed'
        raise e


def _analyze_image_material(material, student):
    """Analyze image material."""
    result = ai_service.analyze_material_image(material.file_path, student.grade, student.id)

    # Save knowledge points
    for kp_data in result.get('knowledge_points', []):
        kp = MaterialKnowledge(
            material_id=material.id,
            subject=material.subject or kp_data.get('subject'),
            chapter=kp_data.get('chapter', ''),
            topic=kp_data.get('topic', ''),
            content=kp_data.get('content', ''),
            key_points=json.dumps(kp_data.get('key_points', []), ensure_ascii=False),
            difficulty=kp_data.get('difficulty', 'medium'),
            importance=kp_data.get('importance', 'important')
        )
        db.session.add(kp)

    # Save questions
    for q_data in result.get('questions', []):
        q = MaterialQuestion(
            material_id=material.id,
            question_text=q_data.get('question_text', ''),
            answer=q_data.get('answer', ''),
            explanation=q_data.get('explanation', ''),
            question_type=q_data.get('question_type', ''),
            knowledge_point=q_data.get('knowledge_point', ''),
            difficulty=q_data.get('difficulty', 'medium'),
            score=q_data.get('score', 5),
            options=json.dumps(q_data.get('options', []), ensure_ascii=False),
            ai_generated=True
        )
        db.session.add(q)

    db.session.commit()


def _analyze_pdf_material(material, student):
    """Analyze PDF material."""
    try:
        import pypdfium2 as pdfium
    except ImportError:
        _analyze_pdf_text_fallback(material, student)
        return

    full_path = Path(__file__).parent.parent / 'app' / 'static' / material.file_path
    pdf = pdfium.PdfDocument(full_path)

    page_images = []
    max_pages = min(len(pdf), 10)

    for page_num in range(max_pages):
        page = pdf[page_num]
        image = page.render(scale=1.0).to_pil()
        image_filename = f"{uuid.uuid4().hex}.jpg"
        image_path = Path(__file__).parent.parent / 'app' / 'static' / 'uploads' / 'materials' / image_filename
        image.save(image_path, 'JPEG', quality=85)
        page_images.append(f"uploads/materials/{image_filename}")

    # Save page images
    material.page_images = json.dumps(page_images, ensure_ascii=False)
    db.session.commit()

    # Analyze first page with full analysis
    first_page_path = Path(__file__).parent.parent / 'app' / 'static' / page_images[0]
    result = ai_service.analyze_material_image(str(page_images[0]), student.grade, student.id)

    # Save knowledge points and questions from first page
    for kp_data in result.get('knowledge_points', []):
        kp = MaterialKnowledge(
            material_id=material.id,
            subject=material.subject or kp_data.get('subject'),
            chapter=kp_data.get('chapter', ''),
            topic=kp_data.get('topic', ''),
            content=kp_data.get('content', ''),
            key_points=json.dumps(kp_data.get('key_points', []), ensure_ascii=False),
            difficulty=kp_data.get('difficulty', 'medium'),
            importance=kp_data.get('importance', 'important')
        )
        db.session.add(kp)

    for q_data in result.get('questions', []):
        q = MaterialQuestion(
            material_id=material.id,
            question_text=q_data.get('question_text', ''),
            answer=q_data.get('answer', ''),
            explanation=q_data.get('explanation', ''),
            question_type=q_data.get('question_type', ''),
            knowledge_point=q_data.get('knowledge_point', ''),
            difficulty=q_data.get('difficulty', 'medium'),
            score=q_data.get('score', 5),
            options=json.dumps(q_data.get('options', []), ensure_ascii=False),
            page_number=1,
            ai_generated=True
        )
        db.session.add(q)

    db.session.commit()

    # Analyze remaining pages
    for page_num in range(1, max_pages):
        page_path = Path(__file__).parent.parent / 'app' / 'static' / page_images[page_num]
        page_result = ai_service.analyze_pdf_page(str(page_images[page_num]), page_num + 1, student.id)

        if page_result.get('is_answer_page'):
            continue

        for q_data in page_result.get('questions', []):
            q = MaterialQuestion(
                material_id=material.id,
                question_text=q_data.get('question_text', ''),
                answer=q_data.get('answer', ''),
                explanation=q_data.get('explanation', ''),
                question_type=q_data.get('question_type', ''),
                knowledge_point=q_data.get('knowledge_point', ''),
                difficulty=q_data.get('difficulty', 'medium'),
                score=q_data.get('score', 5),
                options=json.dumps(q_data.get('options', []), ensure_ascii=False),
                page_number=page_num + 1,
                ai_generated=True
            )
            db.session.add(q)

    db.session.commit()


def _analyze_pdf_text_fallback(material, student):
    """Fallback PDF analysis using text extraction."""
    try:
        import pdfplumber
    except ImportError:
        try:
            import PyPDF2
        except ImportError:
            raise ValueError("No PDF library available")

    full_path = Path(__file__).parent.parent / 'app' / 'static' / material.file_path

    text = ''
    try:
        with pdfplumber.open(full_path) as pdf:
            for page in pdf.pages[:10]:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'
    except Exception:
        with open(full_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages[:10]:
                text += page.extract_text() or ''

    if len(text) > 10000:
        text = text[:10000]

    result = ai_service.analyze_text_content(text, student.grade, student.id)

    for kp_data in result.get('knowledge_points', []):
        kp = MaterialKnowledge(
            material_id=material.id,
            subject=material.subject or kp_data.get('subject'),
            chapter=kp_data.get('chapter', ''),
            topic=kp_data.get('topic', ''),
            content=kp_data.get('content', ''),
            key_points=json.dumps(kp_data.get('key_points', []), ensure_ascii=False),
            difficulty=kp_data.get('difficulty', 'medium'),
            importance=kp_data.get('importance', 'important')
        )
        db.session.add(kp)

    for q_data in result.get('questions', []):
        q = MaterialQuestion(
            material_id=material.id,
            question_text=q_data.get('question_text', ''),
            answer=q_data.get('answer', ''),
            explanation=q_data.get('explanation', ''),
            question_type=q_data.get('question_type', ''),
            knowledge_point=q_data.get('knowledge_point', ''),
            difficulty=q_data.get('difficulty', 'medium'),
            score=q_data.get('score', 5),
            options=json.dumps(q_data.get('options', []), ensure_ascii=False),
            ai_generated=True
        )
        db.session.add(q)

    db.session.commit()


def _analyze_docx_material(material, student):
    """Analyze DOCX material."""
    try:
        from docx import Document
    except ImportError:
        raise ValueError("python-docx not installed")

    full_path = Path(__file__).parent.parent / 'app' / 'static' / material.file_path
    doc = Document(full_path)

    text = '\n'.join([p.text for p in doc.paragraphs])
    if len(text) > 10000:
        text = text[:10000]

    result = ai_service.analyze_text_content(text, student.grade, student.id)

    for kp_data in result.get('knowledge_points', []):
        kp = MaterialKnowledge(
            material_id=material.id,
            subject=material.subject or kp_data.get('subject'),
            chapter=kp_data.get('chapter', ''),
            topic=kp_data.get('topic', ''),
            content=kp_data.get('content', ''),
            key_points=json.dumps(kp_data.get('key_points', []), ensure_ascii=False),
            difficulty=kp_data.get('difficulty', 'medium'),
            importance=kp_data.get('importance', 'important')
        )
        db.session.add(kp)

    for q_data in result.get('questions', []):
        q = MaterialQuestion(
            material_id=material.id,
            question_text=q_data.get('question_text', ''),
            answer=q_data.get('answer', ''),
            explanation=q_data.get('explanation', ''),
            question_type=q_data.get('question_type', ''),
            knowledge_point=q_data.get('knowledge_point', ''),
            difficulty=q_data.get('difficulty', 'medium'),
            score=q_data.get('score', 5),
            options=json.dumps(q_data.get('options', []), ensure_ascii=False),
            ai_generated=True
        )
        db.session.add(q)

    db.session.commit()


def _infer_metadata_from_title(material):
    """Infer metadata from material title."""
    title = material.title

    # Subject
    if not material.subject:
        subject = infer_subject_from_filename(title)
        if subject:
            material.subject = subject

    # Grade
    if not material.source_grade:
        grade = infer_grade_from_filename(title)
        if grade:
            material.source_grade = grade

    # Semester
    if '上册' in title or '第一学期' in title:
        material.source_semester = '上册'
    elif '下册' in title or '第二学期' in title:
        material.source_semester = '下册'

    # Material type
    if not material.material_type:
        if '试卷' in title or '考试' in title:
            material.material_type = '试卷'
        elif '练习' in title:
            material.material_type = '练习册'
        elif '讲义' in title:
            material.material_type = '讲义'
        elif '课本' in title or '教材' in title:
            material.material_type = '课本'
        elif '笔记' in title:
            material.material_type = '笔记'


def _analyze_material_async(material, student):
    """Analyze material in background (simplified sync version)."""
    _analyze_material(material, student)
